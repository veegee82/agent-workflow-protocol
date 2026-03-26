"""Convert paper-awp.md to PDF and DOCX."""

import markdown
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

MD_PATH = Path("paper-awp.md")
PDF_PATH = Path("paper-awp.pdf")
DOCX_PATH = Path("paper-awp.docx")

# ── PDF via WeasyPrint ──────────────────────────────────────────────

def build_pdf():
    md_text = MD_PATH.read_text()
    html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "codehilite"])

    css = CSS(string="""
        @page {
            size: A4;
            margin: 2.5cm 2cm;
            @bottom-center { content: counter(page); font-size: 9pt; color: #666; }
        }
        body {
            font-family: 'Georgia', 'Times New Roman', serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #1a1a1a;
            max-width: 100%;
        }
        h1 {
            font-size: 16pt;
            text-align: center;
            margin-bottom: 0.3em;
            line-height: 1.3;
        }
        h2 { font-size: 13pt; margin-top: 1.5em; border-bottom: 1px solid #ccc; padding-bottom: 4px; }
        h3 { font-size: 11.5pt; margin-top: 1.2em; }
        p { margin: 0.6em 0; text-align: justify; }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
            font-size: 9.5pt;
        }
        th, td {
            border: 1px solid #999;
            padding: 5px 8px;
            text-align: left;
        }
        th { background: #f0f0f0; font-weight: bold; }
        code {
            font-family: 'Courier New', monospace;
            font-size: 9pt;
            background: #f5f5f5;
            padding: 1px 3px;
        }
        pre {
            background: #f5f5f5;
            padding: 10px 12px;
            border: 1px solid #ddd;
            border-radius: 3px;
            font-size: 8.5pt;
            line-height: 1.4;
            overflow-wrap: break-word;
            white-space: pre-wrap;
        }
        pre code { background: none; padding: 0; }
        strong { color: #111; }
        hr { border: none; border-top: 1px solid #ccc; margin: 1.5em 0; }
        em { color: #444; }
    """)

    full_html = f"""<!DOCTYPE html>
    <html><head><meta charset="utf-8"></head>
    <body>{html_body}</body></html>"""

    HTML(string=full_html).write_pdf(str(PDF_PATH), stylesheets=[css])
    print(f"PDF created: {PDF_PATH}")


# ── DOCX via python-docx ───────────────────────────────────────────

def build_docx():
    md_text = MD_PATH.read_text()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    lines = md_text.split("\n")
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        # Filter out separator rows
        data = [r for r in table_rows if not re.match(r"^\|[\s\-:|]+\|$", r)]
        if not data:
            in_table = False
            table_rows = []
            return
        cols = [c.strip() for c in data[0].strip("|").split("|")]
        num_cols = len(cols)
        tbl = doc.add_table(rows=len(data), cols=num_cols, style="Table Grid")
        for i, row_text in enumerate(data):
            cells = [c.strip() for c in row_text.strip("|").split("|")]
            for j in range(min(num_cols, len(cells))):
                cell = tbl.cell(i, j)
                cell.text = cells[j]
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph("")
        in_table = False
        table_rows = []

    for line in lines:
        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Tables
        if line.strip().startswith("|"):
            if not in_table:
                in_table = True
            table_rows.append(line)
            continue
        elif in_table:
            flush_table()

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60)
        elif line.strip() == "":
            continue
        else:
            clean = line.strip()
            p = doc.add_paragraph()
            # Simple bold/italic handling
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", clean)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    flush_table()
    doc.save(str(DOCX_PATH))
    print(f"DOCX created: {DOCX_PATH}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
